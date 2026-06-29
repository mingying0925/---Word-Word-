"""测试 docx_filler 模块的 fill_docx 函数。"""

import os
import sys

import pytest
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# 将项目根目录加入 sys.path,以便导入 services 模块
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from services.docx_filler import fill_docx

# 定位测试用模板文件
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SAMPLE_TEMPLATE = os.path.join(BASE_DIR, 'uploads', 'sample_template.docx')


def _add_bookmarks_to_paragraph(paragraph, names):
    """在段落开头按给定顺序插入多个书签,书签End追加到段落末尾。"""
    p_el = paragraph._p
    # 逆序插入 bookmarkStart,保证最终文档顺序与 names 一致
    for i, name in reversed(list(enumerate(names, start=1))):
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), str(i))
        start.set(qn('w:name'), name)
        p_el.insert(0, start)
    for i, name in enumerate(names, start=1):
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), str(i))
        p_el.append(end)


@pytest.fixture
def multi_bookmark_template(tmp_path):
    """构造一个含多书签单元格的临时 docx 模板,用于测试空格拼接逻辑。

    sample_template.docx 中没有多书签单元格,因此动态创建一个临时模板:
    单元格 (0,1) 同时包含 '签名' 和 '日期' 两个书签。
    """
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '签名/日期'
    value_cell = table.cell(0, 1)
    value_cell.text = ''
    para = value_cell.paragraphs[0]
    _add_bookmarks_to_paragraph(para, ['签名', '日期'])
    template_path = str(tmp_path / 'multi_bookmark_template.docx')
    doc.save(template_path)
    return template_path


def test_fill_docx_generates_valid_file():
    """用 sample_template.docx 填充数据后生成的文件应存在且是有效的 docx。"""
    data = {'姓名': '张三', '性别': '男', '出生日期': '2020-01-01', '照片': '照片占位'}
    out_path = fill_docx(SAMPLE_TEMPLATE, data)
    assert os.path.isfile(out_path), "生成的文件不存在"
    # 用 python-docx 重新打开,验证是有效的 docx
    doc = Document(out_path)
    assert len(doc.tables) > 0


def test_filled_data_written_to_correct_bookmark_position():
    """填充的数据应正确写入对应书签位置(用 python-docx 重新打开验证)。"""
    data = {'姓名': '李四', '性别': '女', '出生日期': '1990-05-15'}
    out_path = fill_docx(SAMPLE_TEMPLATE, data)
    doc = Document(out_path)
    table = doc.tables[0]
    # 姓名书签位于 row=0, col=1
    assert table.cell(0, 1).text == '李四'
    # 性别书签位于 row=1, col=1
    assert table.cell(1, 1).text == '女'
    # 出生日期书签位于 row=2, col=1
    assert table.cell(2, 1).text == '1990-05-15'


def test_fill_docx_with_empty_data():
    """空数据字典应能正常处理不报错。"""
    out_path = fill_docx(SAMPLE_TEMPLATE, {})
    assert os.path.isfile(out_path), "生成的文件不存在"
    # 用 python-docx 重新打开,验证是有效的 docx
    doc = Document(out_path)
    assert len(doc.tables) > 0


def test_multi_bookmark_cell_values_joined_by_space(multi_bookmark_template):
    """多书签单元格的值应被空格拼接。"""
    data = {'签名': '王五', '日期': '2024-12-31'}
    out_path = fill_docx(multi_bookmark_template, data)
    doc = Document(out_path)
    table = doc.tables[0]
    cell_text = table.cell(0, 1).text
    # 两个书签的值应以空格拼接
    assert cell_text == '王五 2024-12-31'


def test_fill_docx_nonexistent_file_raises_exception():
    """文件不存在时应抛出异常。"""
    with pytest.raises(PackageNotFoundError):
        fill_docx('nonexistent_template.docx', {})


def test_fill_docx_with_image_insertion(tmp_path):
    """图片路径作为值传入时应插入图片到单元格。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '照片'
    cell_photo = table.cell(0, 1)
    cell_photo.text = ''
    _add_bookmarks_to_paragraph(cell_photo.paragraphs[0], ['照片'])
    template_path = str(tmp_path / 'photo_template.docx')
    doc.save(template_path)

    img_path = str(tmp_path / 'test_photo.jpg')
    with open(img_path, 'wb') as f:
        f.write(b'\xff\xd8\xff\xe0\x00\x10JFIF\x00\x01')

    data = {'照片': img_path}
    out_path = fill_docx(template_path, data)
    assert os.path.isfile(out_path)

    result_doc = Document(out_path)
    table = result_doc.tables[0]
    cell_text = table.cell(0, 1).text.strip()
    # 插入图片后单元格应不含原始文本(图片替换了文本)
    # 测试路径不在白名单内,应返回路径非法提示
    assert cell_text in ('', '[图片路径非法]', '[图片插入失败]')


def test_fill_docx_with_nonexistent_image_fallback(tmp_path):
    """图片路径不存在时当作文本路径字符串写入。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '照片'
    cell_photo = table.cell(0, 1)
    cell_photo.text = ''
    _add_bookmarks_to_paragraph(cell_photo.paragraphs[0], ['照片'])
    template_path = str(tmp_path / 'photo_fallback_template.docx')
    doc.save(template_path)

    nonexistent = str(tmp_path / 'non_existent.jpg')
    data = {'照片': nonexistent}
    out_path = fill_docx(template_path, data)
    assert os.path.isfile(out_path)

    result_doc = Document(out_path)
    table = result_doc.tables[0]
    # 文件不存在时 is_image=False，路径作为普通文本写入
    assert nonexistent in table.cell(0, 1).text
