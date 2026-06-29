"""测试 Flask 应用的路由接口。"""

import io
import json
import os
import sys

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

import app as app_module


def _add_bookmark_to_paragraph(paragraph, name, bookmark_id=1):
    """在段落中插入一个书签 (bookmarkStart + bookmarkEnd)。"""
    p_el = paragraph._p
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bookmark_id))
    start.set(qn('w:name'), name)
    p_el.insert(0, start)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bookmark_id))
    p_el.append(end)


@pytest.fixture
def client():
    """Flask 测试客户端。"""
    app_module.app.config['TESTING'] = True
    with app_module.app.test_client() as c:
        yield c


@pytest.fixture
def sample_docx(tmp_path):
    """创建含表格与书签的临时 docx 文件。"""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '姓名'
    cell_name = table.cell(0, 1)
    cell_name.text = ''
    _add_bookmark_to_paragraph(cell_name.paragraphs[0], '姓名', bookmark_id=1)
    table.cell(1, 0).text = '性别'
    cell_gender = table.cell(1, 1)
    cell_gender.text = ''
    _add_bookmark_to_paragraph(cell_gender.paragraphs[0], '性别', bookmark_id=2)
    path = str(tmp_path / 'sample.docx')
    doc.save(path)
    return path


def test_health_returns_ok(client):
    """GET /api/health 返回 200。"""
    resp = client.get('/api/health')
    assert resp.status_code == 200
    data = resp.get_json()
    assert data == {"status": "ok", "service": "skillbridge-export"}


def test_parse_bookmarks_no_file_returns_400(client):
    """POST /api/parse-bookmarks 未提供文件时返回 400。"""
    resp = client.post('/api/parse-bookmarks')
    assert resp.status_code == 400


def test_parse_bookmarks_valid_docx_returns_200(client, sample_docx):
    """POST /api/parse-bookmarks 上传有效 docx 时返回 200。"""
    with open(sample_docx, 'rb') as f:
        resp = client.post('/api/parse-bookmarks', data={
            'file': (f, 'sample.docx')
        }, content_type='multipart/form-data')
    assert resp.status_code == 200
    data = resp.get_json()
    assert 'bookmarks' in data
    assert 'tables_structure' in data
    assert len(data['bookmarks']) > 0
    assert len(data['tables_structure']) > 0


def test_fill_word_empty_filename_returns_400(client):
    """POST /api/fill-word 空文件名返回 400。"""
    resp = client.post('/api/fill-word', data={
        'template': (io.BytesIO(b'fake'), ''),
        'data': '{}'
    }, content_type='multipart/form-data')
    assert resp.status_code == 400


def test_fill_word_missing_template_returns_400(client):
    """POST /api/fill-word 未提供模板文件返回 400。"""
    resp = client.post('/api/fill-word', data={
        'data': '{"name":"test"}'
    }, content_type='multipart/form-data')
    assert resp.status_code == 400


def test_fill_word_invalid_json_data_returns_400(client, sample_docx):
    """POST /api/fill-word data 不是合法 JSON 返回 400。"""
    with open(sample_docx, 'rb') as f:
        resp = client.post('/api/fill-word', data={
            'template': (f, 'template.docx'),
            'data': 'not-json'
        }, content_type='multipart/form-data')
    assert resp.status_code == 400


def test_fill_word_with_valid_data_returns_docx(client, sample_docx):
    """POST /api/fill-word 正常填充返回 docx 字节流。"""
    with open(sample_docx, 'rb') as f:
        data = {'姓名': '张三', '性别': '男'}
        resp = client.post('/api/fill-word', data={
            'template': (f, 'template.docx'),
            'data': json.dumps(data)
        }, content_type='multipart/form-data')
    assert resp.status_code == 200
    assert resp.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    # 验证返回的 docx 可打开且内容正确
    result_doc = Document(io.BytesIO(resp.data))
    found_values = []
    for table in result_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found_values.append(cell.text.strip())
    assert '张三' in found_values
    assert '男' in found_values


def test_fill_word_with_empty_data(client, sample_docx):
    """POST /api/fill-word 空 data 也能正常返回 docx。"""
    with open(sample_docx, 'rb') as f:
        resp = client.post('/api/fill-word', data={
            'template': (f, 'template.docx'),
            'data': '{}'
        }, content_type='multipart/form-data')
    assert resp.status_code == 200
    assert resp.content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
