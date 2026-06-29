"""根据书签将数据填充到 Word 模板中,生成填好的 .docx 临时文件。"""

import logging
import os
import tempfile
import uuid

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# 允许的图片目录白名单。
# 优先级：环境变量 ALLOWED_IMAGE_DIRS（多个用 ; 分隔）> 默认相对路径反推。
# Java 端通过启动脚本把 app.upload-dir 注入到 ALLOWED_IMAGE_DIRS，
# 避免 Python 反推路径与实际部署结构不一致导致图片被拒。
_IMAGE_EXTENSIONS = {'.jpg', '.jpeg', '.png', '.gif', '.bmp'}
_MAGIC_BYTES = {
    b'\xff\xd8\xff': 'jpg',
    b'\x89PNG\r\n\x1a\n': 'png',
    b'GIF87a': 'gif',
    b'GIF89a': 'gif',
    b'BM': 'bmp',
}


def _is_image_ext(path):
    """检查文件扩展名是否为允许的图片类型（防止 .jpg.exe 绕过）。"""
    name = os.path.basename(path)
    ext = os.path.splitext(name)[1].lower()
    return ext in _IMAGE_EXTENSIONS


_DEFAULT_ALLOWED = [
    os.path.realpath(os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                  '..', '..', 'skillbridge-web', 'uploads', 'images')),
]


_ALLOWED_DIRS_CACHE = None


def _get_allowed_dirs():
    """返回归一化后的允许图片目录列表（结果缓存）。"""
    global _ALLOWED_DIRS_CACHE
    if _ALLOWED_DIRS_CACHE is not None:
        return _ALLOWED_DIRS_CACHE
    env = os.environ.get('ALLOWED_IMAGE_DIRS')
    if env:
        dirs = [p.strip() for p in env.split(';') if p.strip()]
        _ALLOWED_DIRS_CACHE = [os.path.normcase(os.path.realpath(p)) for p in dirs]
    else:
        _ALLOWED_DIRS_CACHE = [os.path.normcase(p) for p in _DEFAULT_ALLOWED]
    return _ALLOWED_DIRS_CACHE


def _set_paragraph_text(paragraph_el, value_text):
    """将段落内的文本替换为 value_text,尽量保留第一个 run 的样式。

    - 保留第一个 <w:r> 元素及其属性
    - 移除该 run 中所有 <w:t> 子元素,再添加一个新的 <w:t>
    - 移除段落中其他多余的 <w:r>(只保留第一个)
    - 如果段落没有 run,则新建一个
    """
    runs = paragraph_el.findall(qn('w:r'))
    if runs:
        first_run = runs[0]
        for t in first_run.findall(qn('w:t')):
            first_run.remove(t)
        new_t = first_run.makeelement(qn('w:t'))
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = value_text
        first_run.append(new_t)
        for extra in runs[1:]:
            paragraph_el.remove(extra)
    else:
        r = paragraph_el.makeelement(qn('w:r'))
        new_t = r.makeelement(qn('w:t'))
        new_t.set(qn('xml:space'), 'preserve')
        new_t.text = value_text
        r.append(new_t)
        paragraph_el.append(r)


def _clear_paragraph(paragraph_el):
    """清空段落内所有 run 元素,保留段落本身与段落属性。"""
    for r in paragraph_el.findall(qn('w:r')):
        paragraph_el.remove(r)


def _insert_image_into_cell(tc, image_path, doc):
    """将图片插入到表格单元格中，自适应填满单元格宽度。

    清除单元格内现有段落内容，在第一个段落中插入图片。
    图片宽度根据单元格可用宽度自适应，高度按图片原始比例自动缩放。
    """
    paragraphs = tc.findall(qn('w:p'))
    if not paragraphs:
        p = tc.makeelement(qn('w:p'))
        tc.append(p)
        paragraphs = [p]

    # 清除所有段落的现有内容
    for p in paragraphs:
        _clear_paragraph(p)

    from docx.shared import Inches
    from docx.text.paragraph import Paragraph

    # 获取第一个段落的 Paragraph 对象
    para = Paragraph(paragraphs[0], doc)

    # 计算单元格可用宽度
    # tcW 是单元格宽度（单位 dxa = 1/20 pt = 1/1440 inch）
    cell_width_inches = 1.5  # 默认 1.5 英寸
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is not None:
        tcW = tcPr.find(qn('w:tcW'))
        if tcW is not None:
            w_val = tcW.get(qn('w:w'))
            w_type = tcW.get(qn('w:type'))
            if w_val and w_type == 'dxa':
                cell_width_inches = int(w_val) / 1440.0
            elif w_val and w_type == 'pct':
                # 百分比单位（50% = 5000），按表格宽度估算
                cell_width_inches = max(1.0, int(w_val) / 5000.0 * 2.0)

    # 减去单元格内边距（默认左右各 0.08 英寸 = 108 dxa）
    cell_width_inches = max(0.5, cell_width_inches - 0.16)

    # 添加 run 并插入图片，只指定宽度，高度按原始比例自动缩放
    run = para.add_run()
    try:
        run.add_picture(image_path, width=Inches(cell_width_inches))
    except Exception as e:
        logger.warning("图片插入失败: %s, path=%s", e, image_path)
        _set_paragraph_text(paragraphs[0], '[图片插入失败]')


def fill_docx(template_path, data):
    """将 data 字典按书签填入模板,返回生成的临时文件路径。

    - 遍历所有表格的每个单元格
    - 若单元格内含书签名且书签名出现在 data 中,则用 data 值替换单元格文本
    - 若值是字符串且指向已存在的图片文件(.jpg/.png/.jpeg/.gif),则插入真实图片
    """
    doc = Document(template_path)

    for table in doc.tables:
        for row in table.rows:
            for tc in row._tr.findall(qn('w:tc')):
                # 找到本单元格内所有书签（过滤 Word 隐藏书签如 _GoBack）
                bookmark_names = [
                    bm.get(qn('w:name'))
                    for bm in tc.iter(qn('w:bookmarkStart'))
                    if bm.get(qn('w:name')) and not bm.get(qn('w:name')).startswith('_')
                ]
                if not bookmark_names:
                    continue

                # 收集本单元格内所有书签对应的非空值
                matched_values = []
                for name in bookmark_names:
                    if name in data and data[name]:
                        matched_values.append(data[name])
                if not matched_values:
                    continue

                # 判断是否为图片:单值且为存在的图片文件路径
                is_image = (
                    len(matched_values) == 1
                    and isinstance(matched_values[0], str)
                    and os.path.isfile(matched_values[0])
                    and _is_image_ext(matched_values[0])
                )

                # 多书签单元格:分离图片值与文本值,支持图片+文本共存
                image_value = None
                text_values = []
                if not is_image and len(matched_values) > 1:
                    for v in matched_values:
                        if (isinstance(v, str)
                                and os.path.isfile(v)
                                and _is_image_ext(v)):
                            image_value = v
                        else:
                            text_values.append(str(v))
                    if image_value is not None:
                        is_image = True  # 标记为含图片,后续分别处理

                paragraphs = tc.findall(qn('w:p'))
                if not paragraphs:
                    p = tc.makeelement(qn('w:p'))
                    tc.append(p)
                    paragraphs = [p]

                if is_image:
                    # 确定要插入的图片路径
                    img_path = matched_values[0] if image_value is None else image_value
                    real_img_path = os.path.realpath(img_path)
                    # Windows 路径不区分大小写，使用 normcase 统一比较（避免 D: vs d: 不匹配）
                    norm_img_path = os.path.normcase(real_img_path)
                    if not any(norm_img_path.startswith(d + os.sep) for d in _get_allowed_dirs()):
                        logger.warning("图片路径不在允许目录内,已跳过: %s", real_img_path)
                        _set_paragraph_text(paragraphs[0], '[图片路径非法]')
                        continue
                    _insert_image_into_cell(tc, img_path, doc)
                    # 若多书签单元格还有文本值,追加到第二个段落
                    if text_values:
                        if len(paragraphs) < 2:
                            p2 = tc.makeelement(qn('w:p'))
                            tc.append(p2)
                            paragraphs = [paragraphs[0], p2]
                        _set_paragraph_text(paragraphs[1], ' '.join(text_values))
                else:
                    combined = ' '.join(str(v) for v in matched_values)
                    _set_paragraph_text(paragraphs[0], combined)
                    for extra_p in paragraphs[1:]:
                        _clear_paragraph(extra_p)

    temp_dir = tempfile.gettempdir()
    temp_filename = f"filled_{uuid.uuid4().hex}.docx"
    temp_path = os.path.join(temp_dir, temp_filename)
    doc.save(temp_path)
    return temp_path
