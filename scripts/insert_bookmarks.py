#!/usr/bin/env python3
"""
[一次性工具] 在 Word 模板的所有表格单元格中，
根据关键词定位目标单元格并插入书签，保存为新文件。

依赖：python-docx, pywin32（仅 Windows，用于 .doc → .docx 转换）, requests
用法：python insert_bookmarks.py
注意：此脚本为一次性工具，不属于微服务运行时代码。
"""

import os
import sys
import uuid
import tempfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import requests

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
# 模板文件默认从上级目录的 skillbridge-export/uploads/ 查找
TEMPLATE_DIR = os.environ.get('TEMPLATE_DIR', os.path.join(SCRIPT_DIR, '..', 'skillbridge-export', 'uploads'))

# ============================================================
# 1. 书签映射定义
# ============================================================

# (搜索关键词, 书签名称, 方位)
# 方位: 'right' — 该关键词右侧单元格, 'below' — 正下方单元格
SIMPLE_FIELDS = [
    ('姓名',          'name',              'right'),
    ('性别',          'gender',            'right'),
    ('出生年月',      'birth_date',        'right'),
    ('证件类型',      'id_type',           'right'),
    ('证件号码',      'id_number',         'right'),
    ('手机号码',      'phone',             'right'),
    ('考生类别',      'candidate_type',    'right'),
    ('当前最高学历',  'highest_education', 'right'),
    ('学历专业',      'education_major',   'right'),
    ('发证时间',      'education_date',    'right'),
    ('工作单位',      'work_unit',         'right'),
    ('现从事的岗位',  'current_position',  'right'),
    ('申报职业',      'declare_occupation','right'),
    ('申报等级',      'declare_level',     'right'),
    ('考试类型',      'exam_type',         'right'),
    ('申报条件',      'declare_condition', 'right'),
]

CERT_FIELDS = [
    ('已获职业资格',       'current_cert_type', 'right'),
    ('证书等级',           'cert_level',        'right'),
    ('职业（工种）',       'cert_occupation',   'right'),
    ('证书编号',           'cert_number',       'right'),
    ('发证日期',           'cert_date',         'right'),
]

TITLE_FIELDS = [
    ('己获专业技术资格',   'title_type',   'right'),
    ('职称证书',           'title_name',   'right'),
    ('证书等级',           'title_level',  'right'),
    ('证书编号',           'title_number', 'right'),
    ('发证日期',           'title_date',   'right'),
]

EXTRA_FIELDS = [
    ('本人目前从事',   'current_post', 'same_cell'),
    ('累计本职业',     'work_years',   'same_cell'),
    ('照片',           'photo',        'same_cell'),
]

# 工作经历：列索引映射（data row 中的 tc index → 书签前缀）
WORK_COL_MAP = [
    (1, 'work_start'),
    (2, 'work_company'),
    (3, 'work_position'),
    (4, 'work_contact'),
]
WORK_DATA_ROWS = 5

# 以下行范围基于「广东省职业技能等级认定个人申报表」模板结构,
# 若模板行序变更需同步更新。
CERT_ROW_RANGE = (10, 13)    # 证书区域:第 10-12 行
TITLE_ROW_RANGE = (13, 16)   # 职称区域:第 13-15 行

# ============================================================
# 2. 辅助函数
# ============================================================

def _get_cell_text(tc):
    """拼接单元格内所有文本。"""
    parts = []
    for p in tc.findall(qn('w:p')):
        t_nodes = [t.text or '' for t in p.iter(qn('w:t'))]
        parts.append(''.join(t_nodes))
    return ''.join(parts).strip()


def _insert_bookmark(paragraph_el, name, bm_id):
    """在段落中插入 bookmarStart + bookmarkEnd，包裹整个段落。"""
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), str(bm_id))
    start.set(qn('w:name'), name)
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), str(bm_id))
    paragraph_el.insert(0, start)
    paragraph_el.append(end)


def _insert_bookmark_to_cell(tc, name, bm_id):
    """在单元格的第一个段落插入书签。若无可用的段落则新建一个。"""
    paragraphs = tc.findall(qn('w:p'))
    if paragraphs:
        p_el = paragraphs[0]
    else:
        p_el = OxmlElement('w:p')
        tc.append(p_el)
    _insert_bookmark(p_el, name, bm_id)


def _get_tc_index_in_row(tc, row_tr):
    """获取 tc 元素在 <w:tr> 子元素列表中的索引。"""
    tcs = row_tr.findall(qn('w:tc'))
    for idx, child in enumerate(tcs):
        if child is tc:
            return idx
    return -1


def _get_right_cell(tc, row_tr):
    """获取同一行中右侧的下一个 <w:tc>，若不存在返回 None。"""
    tcs = row_tr.findall(qn('w:tc'))
    idx = _get_tc_index_in_row(tc, row_tr)
    if 0 <= idx < len(tcs) - 1:
        return tcs[idx + 1]
    return None


def _get_below_cell(tc, table, row_idx):
    """获取下一行中相同位置的 <w:tc>，若不存在返回 None。"""
    if row_idx + 1 >= len(table.rows):
        return None
    next_row = table.rows[row_idx + 1]
    next_tcs = next_row._tr.findall(qn('w:tc'))
    idx = _get_tc_index_in_row(tc, table.rows[row_idx]._tr)
    if 0 <= idx < len(next_tcs):
        return next_tcs[idx]
    return None


def _find_cell_by_keyword(table, keyword, partial=True, row_range=None):
    """在表格中搜索包含 keyword 的单元格。

    row_range: (start, end) 行号范围（含头不含尾），None 表示搜索全部行。
    返回列表 [(tc, row_tr, row_idx)]
    """
    results = []
    for row_idx, row in enumerate(table.rows):
        if row_range is not None and not (row_range[0] <= row_idx < row_range[1]):
            continue
        tr = row._tr
        for tc in tr.findall(qn('w:tc')):
            text = _get_cell_text(tc)
            if not text:
                continue
            if partial:
                if keyword in text:
                    results.append((tc, tr, row_idx))
            else:
                if text == keyword:
                    results.append((tc, tr, row_idx))
    return results


def _find_cell_by_exact(table, keyword, row_range=None):
    """搜索包含完整关键词的单元格。可指定 row_range 限定行范围。"""
    return _find_cell_by_keyword(table, keyword, partial=True, row_range=row_range)


def _insert_fields_bookmarks(table, fields, bm_id, inserted_count, warnings_list,
                             row_range=None, region_name=''):
    """为字段列表按关键词定位单元格并插入书签。

    返回更新后的 (bm_id, inserted_count)。fields 为 (keyword, bm_name, direction) 列表,
    direction 为 'right' 或 'below'。row_range 限定搜索行范围。
    """
    for keyword, bm_name, direction in fields:
        cells = _find_cell_by_exact(table, keyword, row_range=row_range)
        if not cells:
            scope = f'在{region_name}区域' if region_name else ''
            warnings_list.append(f'关键词 "{keyword}" 未{scope}找到，跳过书签 "{bm_name}"')
            continue

        tc, tr, row_idx = cells[0]

        if direction == 'right':
            target = _get_right_cell(tc, tr)
        else:  # below
            target = _get_below_cell(tc, table, row_idx)

        if target is None:
            warnings_list.append(f'关键词 "{keyword}" 的 {direction} 侧无单元格，跳过书签 "{bm_name}"')
            continue

        _insert_bookmark_to_cell(target, bm_name, bm_id)
        bm_id += 1
        inserted_count += 1
        suffix = f' - {region_name}区域' if region_name else ''
        print(f'  + 书签 "{bm_name}"（关键词 "{keyword}"{suffix}）')

    return bm_id, inserted_count


def convert_doc_to_docx(doc_path):
    """使用 win32com 将 .doc 转换为临时 .docx。"""
    import win32com.client
    word = win32com.client.Dispatch('Word.Application')
    word.Visible = False
    try:
        doc = word.Documents.Open(doc_path)
        temp_path = os.path.join(tempfile.gettempdir(), f'_conv_{uuid.uuid4().hex}.docx')
        doc.SaveAs(temp_path, FileFormat=16)  # wdFormatXMLDocument
        doc.Close()
        return temp_path
    finally:
        word.Quit()


# ============================================================
# 3. 主流程
# ============================================================

def main():
    # --- 找到模板文件 ---
    template_name = None
    search_dir = TEMPLATE_DIR
    if not os.path.isdir(search_dir):
        search_dir = SCRIPT_DIR
    for f in os.listdir(search_dir):
        if f.endswith('.doc') and '申报表' in f:
            template_name = f
            break
    if not template_name:
        # 回退：取任意 .doc 文件
        docs = [f for f in os.listdir(search_dir) if f.endswith('.doc') and not f.endswith('.docx')]
        if not docs:
            print('错误：未找到 .doc 模板文件')
            sys.exit(1)
        template_name = docs[0]

    doc_path = os.path.join(search_dir, template_name)
    print(f'模板文件: {doc_path}')

    # --- 转换 .doc → .docx ---
    try:
        docx_path = convert_doc_to_docx(doc_path)
        print(f'已转换为临时 .docx: {docx_path}')
    except Exception as e:
        print(f'错误：无法转换 .doc 文件，请安装 pywin32 或手动另存为 .docx。{e}')
        sys.exit(1)

    # --- 打开文档 ---
    try:
        document = Document(docx_path)
    except Exception as e:
        print(f'错误：无法打开文档: {e}')
        try:
            os.remove(docx_path)
        except OSError:
            pass
        sys.exit(1)

    if not document.tables:
        print('文档中没有表格，退出。')
        try:
            os.remove(docx_path)
        except OSError:
            pass
        sys.exit(1)

    table = document.tables[0]
    bm_id = 1
    inserted_count = 0
    warnings_list = []

    # ============================================================
    # 3a. 简单字段（SIMPLE_FIELDS）
    # ============================================================
    bm_id, inserted_count = _insert_fields_bookmarks(
        table, SIMPLE_FIELDS, bm_id, inserted_count, warnings_list)

    # ============================================================
    # 3b. 证书区域（CERT_FIELDS）
    # ============================================================
    bm_id, inserted_count = _insert_fields_bookmarks(
        table, CERT_FIELDS, bm_id, inserted_count, warnings_list,
        row_range=CERT_ROW_RANGE, region_name='证书')

    # ============================================================
    # 3c. 职称区域（TITLE_FIELDS）
    # ============================================================
    bm_id, inserted_count = _insert_fields_bookmarks(
        table, TITLE_FIELDS, bm_id, inserted_count, warnings_list,
        row_range=TITLE_ROW_RANGE, region_name='职称')

    # ============================================================
    # 3d. 工作经历区域
    # ============================================================
    # 查找包含 "起止时间" 的行作为 header 行
    header_cells = _find_cell_by_exact(table, '起止时间')
    header_row_idx = None
    if header_cells:
        header_row_idx = header_cells[0][2]  # row_idx
    else:
        warnings_list.append('未找到 "起止时间" 表头，跳过工作经历书签')

    if header_row_idx is not None:
        for i in range(WORK_DATA_ROWS):
            data_row_idx = header_row_idx + 1 + i
            if data_row_idx >= len(table.rows):
                warnings_list.append(f'工作经历第 {i+1} 行超出表格范围，跳过')
                continue

            data_tr = table.rows[data_row_idx]._tr
            data_tcs = data_tr.findall(qn('w:tc'))

            for tc_idx, prefix in WORK_COL_MAP:
                if tc_idx >= len(data_tcs):
                    warnings_list.append(f'工作经历第 {i+1} 行无 tc[{tc_idx}]，跳过书签 "{prefix}_{i+1}"')
                    continue

                bm_name = f'{prefix}_{i+1}'
                target_tc = data_tcs[tc_idx]
                _insert_bookmark_to_cell(target_tc, bm_name, bm_id)
                bm_id += 1
                inserted_count += 1
                print(f'  + 书签 "{bm_name}"（工作经历第 {i+1} 行）')

    # ============================================================
    # 3e. 补充字段（EXTRA_FIELDS）
    # ============================================================
    for keyword, bm_name, location in EXTRA_FIELDS:
        cells = _find_cell_by_exact(table, keyword)
        if not cells:
            warnings_list.append(f'关键词 "{keyword}" 未找到，跳过书签 "{bm_name}"')
            continue

        tc, tr, row_idx = cells[0]
        _insert_bookmark_to_cell(tc, bm_name, bm_id)
        bm_id += 1
        inserted_count += 1
        print(f'  + 书签 "{bm_name}"（关键词 "{keyword}" - 补充字段）')

    # ============================================================
    # 4. 保存文档
    # ============================================================
    output_name = 'template_with_bookmarks.docx'
    output_path = os.path.join(search_dir, output_name)

    try:
        document.save(output_path)
        print(f'\n文件已保存: {output_path}')
    except Exception as e:
        print(f'\n错误：保存文件失败: {e}')
        sys.exit(1)
    finally:
        try:
            os.remove(docx_path)
        except OSError:
            pass

    # ============================================================
    # 5. 输出统计
    # ============================================================
    print(f'\n书签插入完成，共插入 {inserted_count} 个书签')
    if warnings_list:
        print(f'\n警告（共 {len(warnings_list)} 条）:')
        for w in warnings_list:
            print(f'  ⚠ {w}')

    # ============================================================
    # 6. 调用解析接口校验
    # ============================================================
    print('\n正在调用 /api/parse-bookmarks 接口进行校验...')
    try:
        with open(output_path, 'rb') as f:
            resp = requests.post(
                'http://localhost:5000/api/parse-bookmarks',
                files={'file': (output_name, f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
                timeout=30,
            )
        if resp.status_code == 200:
            data = resp.json()
            bookmarks = data.get('bookmarks', [])
            print(f'\n接口返回的书签列表（共 {len(bookmarks)} 个）:')
            for bm in bookmarks:
                bm_name = bm.get('name', '')
                bm_type = bm.get('type', '')
                bm_row = bm.get('row', '')
                bm_col = bm.get('col', '')
                bm_table = bm.get('table_index', '')
                print(f'  [{bm_type:6s}] {bm_name:30s}  table={bm_table} row={bm_row} col={bm_col}')
        else:
            print(f'接口返回错误: HTTP {resp.status_code}')
            print(resp.text[:500])
    except requests.exceptions.ConnectionError:
        print('⚠ 无法连接到本地解析服务（http://localhost:5000），请确保 Flask 服务已启动。')
    except Exception as e:
        print(f'⚠ 调用解析接口时发生异常: {e}')


if __name__ == '__main__':
    main()
